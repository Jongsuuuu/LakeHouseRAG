import os, sys, re, json, hashlib, datetime, unicodedata, warnings
from collections import deque
from dataclasses import dataclass, field
from typing import Dict, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from thefuzz import fuzz
from rank_bm25 import BM25Okapi

warnings.filterwarnings('ignore')

# ── 경로 설정 ─────────────────────────────────────────────────────────────────
API_KEY_PATH = '/content/drive/MyDrive/DATA_LAKEHOUSE/api_key.txt'
DATA_DIR     = '/content/drive/MyDrive/DATA_LAKEHOUSE/DATA'
LAKEHOUSE_DB = '/content/drive/MyDrive/DATA_LAKEHOUSE/structured_lakehouse.parquet'
FAISS_DIR    = '/content/faiss_index_db'
HASH_LOG     = '/content/drive/MyDrive/DATA_LAKEHOUSE/file_hash_log.json'

# ── API 키 로드 ───────────────────────────────────────────────────────────────
def load_api_key():
    with open(API_KEY_PATH, 'r') as f:
        for line in f:
            m = re.match(r'GOOGLE_API_KEY\s*=\s*(.+)', line.strip())
            if m:
                return m.group(1).strip()
    raise ValueError('GOOGLE_API_KEY not found')

os.environ['GOOGLE_API_KEY'] = load_api_key()

# ── 임포트 (무거운 것은 캐시) ─────────────────────────────────────────────────
@st.cache_resource(show_spinner='임베딩 모델 로딩 중...')
def get_embeddings():
    from langchain_huggingface import HuggingFaceEmbeddings
    return HuggingFaceEmbeddings(
        model_name='sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2'
    )

@st.cache_resource(show_spinner='OCR 엔진 초기화 중...')
def get_ocr():
    import easyocr
    return easyocr.Reader(['ko', 'en'], gpu=False)

# ── 텍스트 추출 ───────────────────────────────────────────────────────────────
def md5_of_file(path):
    h = hashlib.md5()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''): h.update(chunk)
    return h.hexdigest()

def extract_pdf_tables(path):
    import pdfplumber
    rows_md = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                for tbl in page.extract_tables() or []:
                    rows = []
                    for j, row in enumerate(tbl):
                        cells = [str(c or '').replace('\n', ' ') for c in row]
                        rows.append('| ' + ' | '.join(cells) + ' |')
                        if j == 0: rows.append('|' + '---|' * len(cells))
                    rows_md.append(f'[표 p{i+1}]\n' + '\n'.join(rows))
    except: pass
    return '\n\n'.join(rows_md)

def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    try:
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f: return f.read()
        elif ext == '.pdf':
            import fitz
            doc = fitz.open(file_path)
            body = '\n'.join(p.get_text() for p in doc)
            tables = extract_pdf_tables(file_path)
            return body + ('\n\n[표]\n' + tables if tables.strip() else '')
        elif ext == '.docx':
            from docx import Document as D
            doc = D(file_path)
            body = '\n'.join(p.text for p in doc.paragraphs)
            tbls = [' | '.join(c.text for c in row.cells) for t in doc.tables for row in t.rows]
            return body + ('\n\n[표]\n' + '\n'.join(tbls) if tbls else '')
        elif ext in ('.xlsx', '.xls'):
            xl = pd.read_excel(file_path, sheet_name=None)
            return '\n\n'.join(f'[{n}]\n{df.to_string(index=False)}' for n, df in xl.items())
        elif ext == '.csv':
            return pd.read_csv(file_path, encoding='utf-8', errors='replace').to_string(index=False)
        elif ext == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), ensure_ascii=False, indent=2)
        elif ext in ('.png','.jpg','.jpeg','.bmp','.tiff','.webp'):
            return ' '.join(get_ocr().readtext(file_path, detail=0, paragraph=True))
    except Exception as e:
        return f'[오류: {e}]'
    return '[지원하지 않는 형식]'

# ── 레이크하우스 동기화 ───────────────────────────────────────────────────────
@st.cache_data(show_spinner='레이크하우스 동기화 중...')
def load_lakehouse(_trigger: int = 0):
    supported = {'.txt','.pdf','.docx','.xlsx','.xls','.csv','.json',
                 '.png','.jpg','.jpeg','.bmp','.tiff','.webp'}
    df = pd.read_parquet(LAKEHOUSE_DB) if os.path.exists(LAKEHOUSE_DB) else pd.DataFrame(
        columns=['file_name','extension','content','path','last_modified','md5'])
    old_hashes = json.load(open(HASH_LOG)) if os.path.exists(HASH_LOG) else {}
    new_hashes = dict(old_hashes)
    new_rows, log = [], []
    for fn in sorted(os.listdir(DATA_DIR)):
        fp = os.path.join(DATA_DIR, fn)
        if not os.path.isfile(fp): continue
        ext = os.path.splitext(fn)[-1].lower()
        if ext not in supported: continue
        h = md5_of_file(fp)
        if old_hashes.get(fn) == h: continue
        df = df[df['file_name'] != fn]
        new_rows.append({'file_name':fn,'extension':ext,'content':extract_text(fp),
                         'path':fp,'last_modified':os.path.getmtime(fp),'md5':h})
        new_hashes[fn] = h
        log.append(fn)
    if new_rows:
        df = pd.concat([df, pd.DataFrame(new_rows)], ignore_index=True)
        df.to_parquet(LAKEHOUSE_DB, index=False)
        json.dump(new_hashes, open(HASH_LOG,'w'), ensure_ascii=False, indent=2)
    return df, log

# ── 벡터 DB ───────────────────────────────────────────────────────────────────
def classify_query(query: str) -> str:
    """AI 기반 쿼리 유형 분류: 'regex' | 'word' | 'sentence'"""
    from langchain_google_genai import ChatGoogleGenerativeAI
    classify_prompt = (
        f'다음 검색 입력의 유형을 판단하세요.\n\n'
        f'입력: "{query}"\n\n'
        '유형 정의:\n'
        '- regex: 정규식 패턴 (특수문자 포함, 패턴 검색 의도)\n'
        '- word: 단어/키워드 검색 (짧은 명사, 고유명사, 용어 등. 공백이 있어도 키워드 나열이면 word)\n'
        '- sentence: 자연어 질문 또는 문장 (의문문, 서술문, 설명/분석 요청 등)\n\n'
        "반드시 'regex', 'word', 'sentence' 중 하나만 답하세요. 다른 말은 절대 하지 마세요."
    )
    try:
        _llm = ChatGoogleGenerativeAI(model='gemini-2.5-flash', temperature=0,
                                       google_api_key=os.environ['GOOGLE_API_KEY'])
        response = _llm.invoke(classify_prompt)
        result = response.content.strip().lower()
        if result in ('regex', 'word', 'sentence'):
            return result
    except Exception:
        pass
    # fallback: LLM 실패 시 규칙 기반
    if any(c in query for c in r'.^$*+?{}[]\\|()'):
        return 'regex'
    if len(query.strip().split()) == 1:
        return 'word'
    return 'sentence'

@st.cache_resource(show_spinner='FAISS 인덱스 빌드 중...')
def get_vectorstore(_trigger: int = 0):
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    df, _ = load_lakehouse(_trigger)
    emb = get_embeddings()
    docs = [Document(page_content=f"--- FILE: {r['file_name']} ---\n{r['content']}",
                     metadata={'file_name':r['file_name'],'extension':r['extension']})
            for _,r in df.iterrows()]
    chunks = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150).split_documents(docs)
    vs = FAISS.from_documents(chunks, emb)
    vs.save_local(FAISS_DIR)
    return vs

# ── 검색 ──────────────────────────────────────────────────────────────────────
def is_hangul(t): return bool(re.search('[가-힣]', t))
def decompose(t): return ''.join(unicodedata.normalize('NFD', t))
def snippet(content, target, w=80):
    cc = content.replace('\n',' ')
    idx = cc.lower().find(target.lower())
    if idx == -1: return cc[:w*2]+'...'
    m = cc[idx:idx+len(target)]
    s,e = max(0,idx-w), min(len(cc),idx+len(m)+w)
    return ('...' if s>0 else '')+cc[s:e].replace(m,f'**{m}**')+('...' if e<len(cc) else '')

def bm25_search(df, query, k=10):
    corpus = [re.sub(r'\s+',' ',r['content']).split() for _,r in df.iterrows()]
    bm = BM25Okapi(corpus)
    scores = bm.get_scores(query.split())
    idx = np.argsort(scores)[::-1][:k]
    return pd.DataFrame([{'파일명':df.iloc[i]['file_name'],'BM25점수':round(float(scores[i]),3),
                           '방식':'BM25','미리보기':snippet(df.iloc[i]['content'],query.split()[0])}
                          for i in idx if scores[i]>0])

def fuzzy_search(df, query, threshold=45):
    results=[]; qc=query.lower().strip(); hg=is_hangul(qc); qp=decompose(qc)
    is_re=any(c in query for c in r'.^$*+?{}[]\|()')
    if is_re:
        try:
            pat=re.compile(query,re.IGNORECASE)
            for _,row in df.iterrows():
                m=pat.search(row['content'])
                if m: results.append({'파일명':row['file_name'],'유사도':100.0,'매칭단어':m.group(),'방식':'Regex','미리보기':snippet(row['content'],m.group())})
            return pd.DataFrame(results)
        except: pass
    for _,row in df.iterrows():
        words=re.sub(r'[^a-zA-Z0-9가-힣\s]',' ',row['content'].lower()).split()
        bs,bw=0,'N/A'
        for w in set(words):
            wp=decompose(w) if hg else w
            if hg and len(qc)>1 and len(w)==1: continue
            if not hg and len(qp)>1 and len(wp)==1: continue
            if len(qp)>=3 and (qp in wp or wp in qp): sc=100.0 if qp==wp else 90.0
            elif hg: sc=(100.0 if qp==wp else fuzz.ratio(qp,wp)); sc*=(0.7 if qc[0]!=w[0] else 1)
            else: sc=fuzz.ratio(qp,wp)*.5+fuzz.token_sort_ratio(qp,wp)*.5; sc*=(0.4 if qc[0]!=w[0] else 1)
            lim=55 if not hg else threshold
            if sc>=lim and sc>bs: bs,bw=sc,w
        if bs>=(55 if not hg else threshold):
            results.append({'파일명':row['file_name'],'유사도':round(float(bs),1),'매칭단어':bw,'방식':'Fuzzy','미리보기':snippet(row['content'],bw)})
    return pd.DataFrame(results).sort_values('유사도',ascending=False) if results else pd.DataFrame()

# ── LLM 답변 ─────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """당신은 최고 수준의 레이크하우스 데이터 분석 전문가입니다.
Chain-of-Thought 사고 과정:
1. [문서 파악] 제공된 문서 목록과 각 파일의 주제를 파악합니다.
2. [근거 추출] 질문과 관련된 핵심 내용을 정확히 찾아냅니다.
3. [종합 분석] 여러 문서 정보를 교차 검증하고 종합합니다.
4. [답변 구성] 명확하고 구조적인 답변을 작성합니다.
5. [출처 명시] 답변에 사용된 파일명을 `📄 출처: [파일명]` 형식으로 표기합니다.

규칙:
- 문서에 없는 내용은 추측하지 마세요.
- 표/수치는 마크다운 표로 정리하세요.
- 한국어로 답변하세요.

검색 힌트: {routing_hint}

관련 문서:
{context}"""

def get_llm_answer(query, history, vs, df, routing_hint, bm25_df, fuzzy_df):
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
    from langchain_classic.chains.combine_documents import create_stuff_documents_chain
    from langchain_core.messages import HumanMessage, AIMessage
    from langchain_core.documents import Document

    llm = ChatGoogleGenerativeAI(model='gemini-2.5-flash', temperature=0.1,
                                  google_api_key=os.environ['GOOGLE_API_KEY'])
    retriever = vs.as_retriever(search_type='mmr', search_kwargs={'k':20,'fetch_k':40})
    chat_history = [HumanMessage(content=m['content']) if m['role']=='user'
                    else AIMessage(content=m['content']) for m in history]

    # BM25/Fuzzy → Document 변환
    extra_docs = []
    for df_, col in [(bm25_df, '파일명'), (fuzzy_df, '파일명')]:
        if df_.empty or col not in df_.columns:
            continue
        for _, row in df_.iterrows():
            matched = df[df['file_name'] == row[col]]
            if matched.empty:
                continue
            extra_docs.append(Document(
                page_content=f"--- FILE: {row[col]} ---\n{row.get('미리보기', '')}",
                metadata={'file_name': row[col], 'source': row.get('방식', 'keyword')}
            ))

    # FAISS 결과
    semantic_docs = retriever.invoke(query)

    # 합치고 중복 제거
    seen = set()
    all_docs = []
    for doc in extra_docs + semantic_docs:
        key = (doc.metadata.get('file_name', ''), doc.page_content[:100])
        if key not in seen:
            seen.add(key)
            all_docs.append(doc)

    prompt = ChatPromptTemplate.from_messages([
        ('system', SYSTEM_PROMPT),
        MessagesPlaceholder(variable_name='chat_history'),
        ('human', '{input}'),
    ])
    doc_chain = create_stuff_documents_chain(llm, prompt)
    result = doc_chain.invoke({
        'input': query,
        'context': all_docs,
        'chat_history': chat_history,
        'routing_hint': routing_hint,
    })

    sources = list({doc.metadata.get('file_name', '') for doc in all_docs})
    return result, sources
# ── Streamlit UI ──────────────────────────────────────────────────────────────
st.set_page_config(page_title='LakehouseRAG', page_icon='🏛️', layout='wide')

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700&family=Space+Mono:wght@700&display=swap');
html, body, [class*="css"] { font-family: 'Noto Sans KR', sans-serif; }
h1 { font-family: 'Space Mono', monospace; }
.stTabs [data-baseweb="tab"] { font-size: 1rem; font-weight: 700; }
.chat-user { background:#dbeafe; border-left:4px solid #3b82f6;
             padding:12px 16px; border-radius:8px; margin:8px 0; }
.chat-bot  { background:#dcfce7; border-left:4px solid #22c55e;
             padding:12px 16px; border-radius:8px; margin:8px 0; }
.source-badge { background:#e0e7ff; color:#3730a3; padding:2px 10px;
                border-radius:12px; font-size:0.8rem; margin:2px; display:inline-block; }
</style>
""", unsafe_allow_html=True)

st.title('🏛️ LakehouseRAG')
st.caption('BM25 + Fuzzy + Semantic 3중 하이브리드 지능형 검색 시스템')

# 세션 상태 초기화
if 'chat_history' not in st.session_state: st.session_state.chat_history = []
if 'reload_trigger' not in st.session_state: st.session_state.reload_trigger = 0

df, updated_files = load_lakehouse(st.session_state.reload_trigger)
vs = get_vectorstore(st.session_state.reload_trigger)

# ── 탭 ────────────────────────────────────────────────────────────────────────
tab_chat, tab_files, tab_dash = st.tabs(['💬 채팅 검색', '📁 파일 탐색기', '📊 현황 대시보드'])

# ── 탭 1: 채팅 ────────────────────────────────────────────────────────────────
with tab_chat:
    col_chat, col_sidebar = st.columns([3, 1])
    with col_sidebar:
        st.markdown('**⚙️ 설정**')
        if st.button('🔄 데이터 새로고침', use_container_width=True):
            st.session_state.reload_trigger += 1
            st.cache_data.clear(); st.cache_resource.clear()
            st.rerun()
        if st.button('🗑️ 대화 초기화', use_container_width=True):
            st.session_state.chat_history = []
            st.rerun()
        st.divider()
        st.markdown(f'**📂 적재 파일 수:** {len(df)}')
        st.markdown(f'**💾 DB 크기:** {os.path.getsize(LAKEHOUSE_DB)/1024:.1f} KB' if os.path.exists(LAKEHOUSE_DB) else '')

    with col_chat:
        # 채팅 히스토리 렌더링
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.markdown(f'<div class="chat-user">👤 <b>사용자</b><br>{msg["content"]}</div>',
                            unsafe_allow_html=True)
            else:
                qtype = msg.get('query_type', 'sentence')
                if qtype in ('word', 'regex'):
                    label = '🔍 Fuzzy 검색 결과' if qtype == 'word' else '🔎 Regex 검색 결과'
                    if msg.get('fuzzy') is not None and not msg['fuzzy'].empty:
                        with st.expander(label, expanded=False):
                            st.dataframe(msg['fuzzy'], use_container_width=True)
                st.markdown(f'<div class="chat-bot">🤖 <b>AI 분석 답변</b></div>',
                            unsafe_allow_html=True)
                st.markdown(msg['content'])
                if msg.get('sources'):
                    badges = ' '.join(f'<span class="source-badge">📄 {s}</span>' for s in msg['sources'])
                    st.markdown(f'참조: {badges}', unsafe_allow_html=True)
                st.divider()

        # 입력
        query = st.chat_input('검색어 또는 질문을 입력하세요...')
        if query:
            st.session_state.chat_history.append({'role':'user','content':query})
            # 슬라이딩 윈도우 (최근 12 메시지)
            window = st.session_state.chat_history[-12:]
            hist_for_llm = [m for m in window if m['role'] in ('user','assistant')]
            # 하이브리드 검색
            qtype = classify_query(query)
            bm25_df = pd.DataFrame()
            fuzzy_df = pd.DataFrame()

            if qtype == 'regex':
                fuzzy_df = fuzzy_search(df, query)
                hint_files = set(fuzzy_df['파일명'].unique()) if not fuzzy_df.empty else set()
                routing_hint = f"정규식 검색 관련 파일: {', '.join(hint_files)}" if hint_files else '정규식 매칭 없음.'
            elif qtype == 'word':
                fuzzy_df = fuzzy_search(df, query)
                hint_files = set(fuzzy_df['파일명'].unique()) if not fuzzy_df.empty else set()
                routing_hint = f"단어 Fuzzy 검색 관련 파일: {', '.join(hint_files)}" if hint_files else '단어 매칭 없음.'
            else:  # sentence
                routing_hint = '문장 질문. 의미 기반(Semantic) 검색만 활용.'

            with st.spinner('🤖 AI가 분석 중입니다...'):
                answer, sources = get_llm_answer(query, hist_for_llm, vs, df, routing_hint, bm25_df, fuzzy_df)
            st.session_state.chat_history.append({
                'role':'assistant','content':answer,'sources':sources,
                'bm25':bm25_df,'fuzzy':fuzzy_df,'query_type':qtype
            })
            st.rerun()

# ── 탭 2: 파일 탐색기 ─────────────────────────────────────────────────────────
with tab_files:
    st.subheader('📁 레이크하우스 파일 탐색기')
    if df.empty:
        st.info('데이터가 없습니다. 데이터 디렉토리를 확인하세요.')
    else:
        ext_filter = st.multiselect('확장자 필터', options=sorted(df['extension'].unique()),
                                     default=sorted(df['extension'].unique()))
        search_filter = st.text_input('파일명 검색')
        filtered = df[df['extension'].isin(ext_filter)]
        if search_filter:
            filtered = filtered[filtered['file_name'].str.contains(search_filter, case=False, na=False)]
        display_df = filtered[['file_name','extension','last_modified','md5']].copy()
        display_df['last_modified'] = pd.to_datetime(display_df['last_modified'], unit='s').dt.strftime('%Y-%m-%d %H:%M')
        display_df.columns = ['파일명','확장자','최종수정','MD5']
        st.dataframe(display_df, use_container_width=True, height=400)

        st.divider()
        selected_file = st.selectbox('파일 내용 미리보기', options=filtered['file_name'].tolist())
        if selected_file:
            content = filtered[filtered['file_name']==selected_file].iloc[0]['content']
            st.text_area('내용 (처음 3000자)', value=content[:3000], height=300)

# ── 탭 3: 대시보드 ────────────────────────────────────────────────────────────
with tab_dash:
    st.subheader('📊 레이크하우스 현황 대시보드')
    if df.empty:
        st.info('데이터가 없습니다.')
    else:
        c1, c2, c3, c4 = st.columns(4)
        c1.metric('📁 총 파일 수', len(df))
        c2.metric('📝 총 문자 수', f"{df['content'].str.len().sum():,}")
        c3.metric('💾 DB 크기', f"{os.path.getsize(LAKEHOUSE_DB)/1024:.1f} KB" if os.path.exists(LAKEHOUSE_DB) else '-')
        c4.metric('🔄 이번 실행 업데이트', len(updated_files))

        st.divider()
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown('**확장자별 파일 수**')
            ext_cnt = df['extension'].value_counts().reset_index()
            ext_cnt.columns = ['확장자','파일 수']
            st.bar_chart(ext_cnt.set_index('확장자'))
        with col_b:
            st.markdown('**파일별 텍스트 길이 (Top 15)**')
            top15 = df.assign(길이=df['content'].str.len()).nlargest(15,'길이')[['file_name','길이']]
            top15.columns = ['파일명','텍스트 길이']
            st.bar_chart(top15.set_index('파일명'))

        if updated_files:
            st.divider()
            st.markdown('**이번 실행에서 업데이트된 파일:**')
            for f_name in updated_files:
                st.markdown(f'- 🔄 `{f_name}`')
